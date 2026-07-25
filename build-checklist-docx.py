# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn

path = r"C:\Users\marce\OneDrive\Área de Trabalho\despachante-sergio-neto\checklist-google-meu-negocio.docx"
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def add_heading_custom(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x33, 0x58)
    return h


def add_check(text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run("☐  ")
    run.bold = True
    p.add_run(text)
    return p


def add_para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_box_label(label, content):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x33, 0x58)
    doc.add_paragraph(content)


title = doc.add_heading(
    "Despachante Sérgio Neto — Checklist Google Meu Negócio + Posts",
    level=0,
)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x33, 0x58)

add_para(
    "Use as caixas ☐ para ir marcando. Quando terminar um item, troque por ☑.",
    italic=True,
    size=10,
)
add_para("Site: https://marcelofigueiredo23.github.io/despachante-sergio-neto/", size=10)
add_para("Instagram: @despachante_sergioneto", size=10)
add_para("Painel: https://business.google.com", size=10)
doc.add_paragraph()

add_heading_custom("0. Dados oficiais (padronizar em tudo)", 1)
add_para("Use exatamente estes dados no Google, Instagram, site e WhatsApp:")
items = [
    ("Nome do negócio: ", "Despachante Sérgio Neto"),
    ("Apelido (só na descrição): ", "Biju"),
    ("Endereço: ", "Av. Éster, 119 — Bela Vista III, Cosmópolis-SP, CEP 13150-000"),
    ("Telefone: ", "(19) 3872-1151"),
    ("Site: ", "https://marcelofigueiredo23.github.io/despachante-sergio-neto/"),
    ("Instagram: ", "https://www.instagram.com/despachante_sergioneto/"),
    (
        "Maps: ",
        "https://www.google.com/maps/place/Sergio+Biju+despachante/@-22.6485522,-47.1855965,17z",
    ),
]
for label, val in items:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(label)
    r.bold = True
    p.add_run(val)

doc.add_paragraph()
add_para(
    "WhatsApp (confirmar número celular): _______________________________",
    italic=True,
)

add_heading_custom("1. Checklist — configurar o perfil (faça nesta ordem)", 1)

add_heading_custom("1.1 Acesso e nome", 2)
for t in [
    "Abrir https://business.google.com e entrar com a conta Google do negócio",
    "Reivindicar / confirmar que é o gestor do perfil atual",
    "Alterar o nome para: Despachante Sérgio Neto",
    "Categoria principal: Despachante (ou a mais próxima disponível)",
    "Adicionar categorias secundárias (licenciamento / documentação veicular, se existir)",
]:
    add_check(t)

add_heading_custom("1.2 Dados de contato e localização", 2)
for t in [
    "Conferir endereço: Av. Éster, 119 — Bela Vista III, Cosmópolis-SP, 13150-000",
    "Marcar o pin no mapa na porta correta do escritório",
    "Telefone: (19) 3872-1151",
    "Ativar WhatsApp com o número celular correto",
    "Colar o site do GitHub no campo Website",
    "Adicionar link do Instagram (se houver campo de redes)",
    "Definir área de atendimento (Cosmópolis + cidades próximas, se atender)",
]:
    add_check(t)

add_heading_custom("1.3 Horário de funcionamento", 2)
add_para("Preencha o horário real. Exemplo (ajuste se for diferente):")
add_para("Segunda a Sexta: 08:00 – 18:00")
add_para("Sábado: ____  |  Domingo: Fechado")
for t in [
    "Cadastrar horário de todos os dias da semana",
    "Marcar feriados / horário especial quando houver",
]:
    add_check(t)

add_heading_custom("1.4 Descrição do negócio (copiar e colar)", 2)
add_check("Colar a descrição abaixo no perfil")
doc.add_paragraph(
    "Despachante em Cosmópolis-SP. Atendimento próximo e ágil para transferência de veículo, "
    "licenciamento, IPVA, primeiro emplacamento, segunda via e documentação veicular. "
    "Conhecido na cidade como Biju. Av. Éster, 119 — Bela Vista III. "
    "Atendimento por telefone e WhatsApp."
)

add_heading_custom("1.5 Fotos (meta: 15–25)", 2)
for t in [
    "Foto de capa / logo",
    "Fachada com número da porta",
    "Interior / balcão de atendimento",
    "Equipe (1–2 fotos)",
    "Mais 6–10 fotos do escritório (luz do dia)",
    "Não publicar documento de cliente nem selfie escura",
]:
    add_check(t)

add_heading_custom("1.6 Serviços (cadastrar os 6 abaixo)", 2)
for t in [
    "Transferência de veículo",
    "Licenciamento anual",
    "IPVA e débitos do veículo",
    "Primeiro emplacamento",
    "Segunda via de documentos",
    "Alteração de endereço e características",
]:
    add_check(t)

add_heading_custom("1.7 Avaliações", 2)
for t in [
    "Pedir avaliação para 5 clientes desta semana",
    "Meta: 15–20 avaliações em 60 dias",
    "Responder todas as avaliações em até 48h",
]:
    add_check(t)

add_heading_custom("1.8 Perguntas e respostas (criar e responder)", 2)
qas = [
    (
        "Quanto custa a transferência?",
        "Depende do veículo e da documentação. Peça orçamento pelo WhatsApp ou telefone (19) 3872-1151.",
    ),
    (
        "Vocês atendem sábado?",
        "Consulte nosso horário no Google ou chame no WhatsApp para confirmar.",
    ),
    (
        "Onde fica o despachante?",
        "Av. Éster, 119 — Bela Vista III, Cosmópolis-SP.",
    ),
    (
        "Fazem licenciamento e IPVA?",
        "Sim. Cuidamos de licenciamento, IPVA, débitos e documentação veicular.",
    ),
]
for q, a in qas:
    add_check(f"Publicar Q: {q}")
    p = doc.add_paragraph()
    r = p.add_run("Resposta: ")
    r.bold = True
    p.add_run(a)

add_heading_custom("2. Textos prontos dos 6 serviços (colar no Google)", 1)

services = [
    (
        "Transferência de veículo",
        "Cuidamos da transferência de propriedade no DETRAN: orientação de documentos, acompanhamento do processo e regularização para o novo proprietário circular em dia. Atendimento em Cosmópolis-SP.",
        "Sob consulta (WhatsApp)",
    ),
    (
        "Licenciamento anual",
        "Emissão e acompanhamento do licenciamento do veículo. Você traz os dados do carro e a gente resolve a documentação para você não perder prazo nem circular irregular.",
        "Sob consulta (WhatsApp)",
    ),
    (
        "IPVA e débitos do veículo",
        "Consulta e orientação para quitar IPVA, multas e débitos vinculados ao veículo. Organizamos o que precisa ser pago e o que pode ser resolvido na documentação.",
        "Sob consulta (WhatsApp)",
    ),
    (
        "Primeiro emplacamento",
        "Documentação para veículo zero km ou situações de primeiro emplacamento. Orientamos a lista correta de documentos e acompanhamos o trâmite até a liberação.",
        "Sob consulta (WhatsApp)",
    ),
    (
        "Segunda via de documentos",
        "Segunda via de CRLV e demais documentos veiculares. Ideal para quem perdeu, danificou ou precisa regularizar a documentação com agilidade.",
        "Sob consulta (WhatsApp)",
    ),
    (
        "Alteração de endereço e características",
        "Atualização de endereço no documento do veículo e alterações de características quando necessário. Evite problemas em blitz e na hora de vender o carro.",
        "Sob consulta (WhatsApp)",
    ),
]
for i, (name, desc, price) in enumerate(services, 1):
    add_heading_custom(f"2.{i} {name}", 2)
    add_check(f"Cadastrar serviço no Google: {name}")
    add_box_label("Descrição:", desc)
    add_box_label("Preço:", price)

add_heading_custom("3. Calendário de posts (Google + Instagram)", 1)
add_para(
    "Para cada post: anexe o criativo logo abaixo da copy, marque o ☐ quando publicar, e anote a data.",
    italic=True,
    size=10,
)

posts = [
    (
        "POST 1 — Site no ar (publicar primeiro)",
        "Hoje / assim que possível",
        "Google Meu Negócio + Instagram",
        "Criativo: arte do site / “Agora temos site”",
        (
            "Agora temos site!\n"
            "Conheça nossos serviços e fale direto pelo WhatsApp:\n"
            "🌐 https://marcelofigueiredo23.github.io/despachante-sergio-neto/\n"
            "Instagram: @despachante_sergioneto\n"
            "Despachante em Cosmópolis — transferência, licenciamento, IPVA e mais."
        ),
    ),
    (
        "POST 2 — Licenciamento",
        "Daqui 3–4 dias",
        "Google Meu Negócio + Instagram",
        "Criativo: arte de licenciamento",
        (
            "Já pensou no licenciamento do seu veículo?\n"
            "Não deixe para a última hora. No Despachante Sérgio Neto (Biju) a gente cuida da documentação com atendimento próximo em Cosmópolis.\n"
            "📍 Av. Éster, 119 — Bela Vista III\n"
            "📲 Chama no WhatsApp e tire sua dúvida."
        ),
    ),
    (
        "POST 3 — Transferência",
        "Na semana seguinte",
        "Google Meu Negócio + Instagram",
        "Criativo: arte de transferência / compra e venda",
        (
            "Comprou ou vendeu um carro?\n"
            "A transferência precisa estar correta para não dar dor de cabeça depois.\n"
            "Aqui orientamos os documentos certos e acompanhamos o processo.\n"
            "Despachante Sérgio Neto — Cosmópolis-SP\n"
            "Fale conosco e resolva sem enrolação."
        ),
    ),
    (
        "POST 4 — Documentos (educativo)",
        "Na semana seguinte ao Post 3",
        "Google Meu Negócio + Instagram",
        "Criativo: lista de documentos / checklist",
        (
            "Documentos mais pedidos na transferência:\n"
            "✅ Documento do comprador e do vendedor\n"
            "✅ CRLV / documentação do veículo\n"
            "✅ Comprovante de endereço\n"
            "✅ Recibo de compra e venda preenchido\n"
            "Lista pode variar — confirme conosco antes de ir ao cartório/DETRAN.\n"
            "Despachante Sérgio Neto | Cosmópolis"
        ),
    ),
]

for title_p, when, where, creative, copy in posts:
    add_heading_custom(title_p, 2)
    add_check("Publicar — data: __________")
    p = doc.add_paragraph()
    r = p.add_run("Quando: ")
    r.bold = True
    p.add_run(when)
    p = doc.add_paragraph()
    r = p.add_run("Onde: ")
    r.bold = True
    p.add_run(where)
    p = doc.add_paragraph()
    r = p.add_run("Criativo: ")
    r.bold = True
    p.add_run(creative)
    add_para("⬇️ Cole / anexe a imagem do criativo aqui embaixo:", italic=True, size=10)
    doc.add_paragraph("[ ESPAÇO PARA O CRIATIVO ]")
    add_box_label("Copy pronta (copiar e colar):", copy)
    add_check("Publicado no Google")
    add_check("Publicado no Instagram")
    add_check("Stories / reforço (opcional)")
    doc.add_paragraph("— — —")

add_heading_custom("4. Texto pronto para pedir avaliação", 1)
add_para(
    "Se puder me ajudar com uma avaliação no Google, me ajuda bastante a aparecer para mais gente da cidade 🙏\n"
    "É só entrar no perfil e deixar sua estrela. Qualquer coisa, estamos à disposição!"
)
add_check("Enviar para clientes após atendimento ok")

add_heading_custom("5. Checklist da Semana 1 (1 hora + follow-up)", 1)
for t in [
    "Dia 1: nome, categoria, endereço, telefone, site, horário, descrição",
    "Dia 1: subir pelo menos 10 fotos",
    "Dia 1: cadastrar os 6 serviços",
    "Dia 1: publicar POST 1 (site) no Google e Instagram",
    "Dia 2–3: criar as 4 perguntas e respostas",
    "Dia 3–4: publicar POST 2 (licenciamento)",
    "Durante a semana: pedir avaliação a 5 clientes",
    "Semana 2: POST 3 + POST 4",
]:
    add_check(t)

add_heading_custom("6. Observações", 1)
add_para("• Não compre avaliações — o Google pode derrubar o perfil.")
add_para("• Nome, telefone e endereço precisam ser iguais no Google, site e Instagram.")
add_para('• “Biju” fica na descrição e nos posts, não no nome oficial do perfil.')

doc.save(path)
print("OK", path)
