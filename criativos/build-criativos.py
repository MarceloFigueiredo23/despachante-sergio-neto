# -*- coding: utf-8 -*-
"""Gera 4 criativos 1080x1080 com foto de contexto + DOCX."""
from __future__ import annotations

import subprocess
import time
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = Path(__file__).resolve().parent
BG = OUT / "bg"
LOGO = ROOT / "logo.png"
EDGE = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")

WA_LINE = "19 99185-1849 · 19 99791-1339"
ADDR = "Campo do Kigol · Novo Horizonte, Cosmópolis-SP"
SITE = "https://despachantesergioneto.com.br/"

POSTS = [
    {
        "slug": "01-site",
        "bg": "01-site.jpg.png",
        "dia": "Dia 1 — publicar hoje",
        "eyebrow": "COSMÓPOLIS · SP",
        "title": "AGORA TEMOS<br>SITE",
        "body": "Seu despachante presencial, com a melhor agilidade online. Lojista ou transportadora? Condições especiais.",
        "cta": WA_LINE,
        "copy": (
            "Agora temos site!\n"
            "O seu despachante presencial, com a melhor agilidade online.\n"
            f"🌐 {SITE}\n"
            "📲 WhatsApp: (19) 99185-1849 · (19) 99791-1339\n"
            "📸 Instagram: @despachante_sergioneto\n"
            f"📍 {ADDR}\n"
            "Lojistas e transportadoras: condições especiais."
        ),
    },
    {
        "slug": "02-licenciamento",
        "bg": "02-licenciamento.jpg.png",
        "dia": "Dia 4 — daqui 3–4 dias",
        "eyebrow": "NÃO DEIXE PARA DEPOIS",
        "title": "LICENCIAMENTO<br>EM DIA",
        "body": "Circular tranquilo começa com a documentação certa. Atendimento presencial + agilidade online.",
        "cta": WA_LINE,
        "copy": (
            "Já pensou no licenciamento do seu veículo?\n"
            "Não deixe para a última hora. No Despachante Sérgio Neto (Biju) a gente resolve com atendimento próximo.\n"
            f"📍 {ADDR}\n"
            "📲 WhatsApp: (19) 99185-1849 · (19) 99791-1339\n"
            "📸 @despachante_sergioneto"
        ),
    },
    {
        "slug": "03-transferencia",
        "bg": "03-transferencia.jpg.png",
        "dia": "Semana 2",
        "eyebrow": "COMPROU OU VENDEU?",
        "title": "TRANSFERÊNCIA<br>SEM DOR DE CABEÇA",
        "body": "Negócio fechado só fica bom com a transferência certa. Orientamos os documentos e acompanhamos o processo.",
        "cta": WA_LINE,
        "copy": (
            "Comprou ou vendeu um carro?\n"
            "A transferência precisa estar correta para não dar dor de cabeça depois.\n"
            "Despachante Sérgio Neto — Cosmópolis-SP\n"
            f"📍 {ADDR}\n"
            "📲 WhatsApp: (19) 99185-1849 · (19) 99791-1339\n"
            "📸 @despachante_sergioneto"
        ),
    },
    {
        "slug": "04-documentos",
        "bg": "04-documentos.jpg.png",
        "dia": "Semana 2 — após o post 3",
        "eyebrow": "CHECKLIST RÁPIDO",
        "title": "DOCUMENTOS DA<br>TRANSFERÊNCIA",
        "body": "RG/CNH · CRLV · Comprovante de endereço · Recibo. Confirme a lista conosco antes de sair de casa.",
        "cta": WA_LINE,
        "copy": (
            "Documentos mais pedidos na transferência:\n"
            "✅ Documento do comprador e do vendedor\n"
            "✅ CRLV / documentação do veículo\n"
            "✅ Comprovante de endereço\n"
            "✅ Recibo de compra e venda preenchido\n"
            "Lista pode variar — confirme conosco antes.\n"
            f"📍 {ADDR}\n"
            "📲 WhatsApp: (19) 99185-1849 · (19) 99791-1339\n"
            "📸 @despachante_sergioneto"
        ),
    },
]


def html_for(post: dict) -> str:
    logo_uri = LOGO.as_uri()
    bg_uri = (BG / post["bg"]).as_uri()
    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8" />
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  html, body {{ width: 1080px; height: 1080px; overflow: hidden; }}
  body {{
    font-family: "Segoe UI", Arial, sans-serif;
    color: #eef3f9;
    background: #0d1a2b url("{bg_uri}") center/cover no-repeat;
    position: relative;
  }}
  .scrim {{
    position: absolute; inset: 0;
    background:
      linear-gradient(115deg, rgba(10,22,40,.92) 0%, rgba(10,22,40,.78) 42%, rgba(10,22,40,.35) 68%, rgba(10,22,40,.55) 100%),
      linear-gradient(180deg, rgba(10,22,40,.25) 0%, transparent 35%, rgba(10,22,40,.88) 100%);
  }}
  .frame {{
    position: relative; z-index: 1;
    width: 1080px; height: 1080px;
    padding: 56px 60px 52px;
    display: flex; flex-direction: column;
  }}
  .top {{
    display: flex; align-items: center; gap: 20px;
    margin-bottom: 56px;
  }}
  .logo {{
    width: 132px; height: 132px; border-radius: 50%;
    background: #fff;
    box-shadow: 0 0 0 4px rgba(232,163,23,.7), 0 16px 36px rgba(0,0,0,.45);
  }}
  .brand-line {{
    font-size: 20px; font-weight: 700; letter-spacing: .14em;
    text-transform: uppercase; color: rgba(238,243,249,.82);
    text-shadow: 0 2px 12px rgba(0,0,0,.45);
  }}
  .eyebrow {{
    display: inline-flex; align-items: center; gap: 12px;
    font-size: 20px; font-weight: 800; letter-spacing: .16em;
    color: #f0b429; margin-bottom: 18px;
    text-shadow: 0 2px 10px rgba(0,0,0,.5);
  }}
  .eyebrow::before {{ content: ""; width: 36px; height: 3px; background: #f0b429; }}
  h1 {{
    font-family: Impact, "Arial Black", sans-serif;
    font-size: 84px; line-height: .94; letter-spacing: .02em;
    font-weight: 400; margin-bottom: 22px; max-width: 11ch;
    text-transform: uppercase;
    color: #ffffff;
    -webkit-text-fill-color: #ffffff;
    text-shadow: 0 4px 22px rgba(0,0,0,.55);
  }}
  .body {{
    font-size: 28px; line-height: 1.38; font-weight: 600;
    color: rgba(238,243,249,.94); max-width: 17ch; margin-bottom: auto;
    text-shadow: 0 2px 14px rgba(0,0,0,.55);
  }}
  .footer {{
    display: flex; align-items: center; justify-content: space-between;
    gap: 16px;
    border-top: 1px solid rgba(238,243,249,.22);
    padding-top: 24px; margin-top: 28px;
  }}
  .cta {{
    display: inline-flex; align-items: center; gap: 12px;
    background: #25d366; color: #062312;
    font-weight: 800; font-size: 20px;
    padding: 14px 20px; border-radius: 999px;
    box-shadow: 0 10px 24px rgba(0,0,0,.35);
    max-width: 640px;
  }}
  .wa {{
    width: 36px; height: 36px; flex-shrink: 0;
    display: block;
  }}
  .handle {{
    text-align: right; font-size: 17px; font-weight: 700;
    color: rgba(238,243,249,.8); letter-spacing: .03em;
    text-shadow: 0 2px 10px rgba(0,0,0,.5);
    line-height: 1.35;
  }}
  .handle small {{
    display: block; font-size: 14px; font-weight: 600;
    color: rgba(238,243,249,.55); margin-top: 4px;
  }}
</style>
</head>
<body>
  <div class="scrim"></div>
  <div class="frame">
    <div class="top">
      <img class="logo" src="{logo_uri}" alt="" />
      <div class="brand-line">Despachante Sérgio Neto · Biju</div>
    </div>
    <div class="eyebrow">{post['eyebrow']}</div>
    <h1>{post['title']}</h1>
    <p class="body">{post['body']}</p>
    <div class="footer">
      <div class="cta">
        <svg class="wa" viewBox="0 0 24 24" aria-hidden="true" xmlns="http://www.w3.org/2000/svg">
          <path fill="#ffffff" d="M17.472 14.382c-.297-.149-1.758-.867-2.03-.967-.273-.099-.471-.148-.67.15-.197.297-.767.966-.94 1.164-.173.199-.347.223-.644.075-.297-.15-1.255-.463-2.39-1.475-.883-.788-1.48-1.761-1.653-2.059-.173-.297-.018-.458.13-.606.134-.133.298-.347.446-.52.149-.174.198-.298.298-.497.099-.198.05-.371-.025-.52-.075-.149-.669-1.612-.916-2.207-.242-.579-.487-.5-.669-.51-.173-.008-.371-.01-.57-.01-.198 0-.52.074-.792.372-.272.297-1.04 1.016-1.04 2.479 0 1.462 1.065 2.875 1.213 3.074.149.198 2.096 3.2 5.077 4.487.709.306 1.262.489 1.694.625.712.227 1.36.195 1.871.118.571-.085 1.758-.719 2.006-1.413.248-.694.248-1.289.173-1.413-.074-.124-.272-.198-.57-.347m-5.421 7.403h-.004a9.87 9.87 0 01-5.031-1.378l-.361-.214-3.741.982.998-3.648-.235-.374a9.86 9.86 0 01-1.51-5.26c.001-5.45 4.436-9.884 9.888-9.884 2.64 0 5.122 1.03 6.988 2.898a9.825 9.825 0 012.893 6.994c-.003 5.45-4.435 9.884-9.885 9.884m8.413-18.297A11.815 11.815 0 0012.05 0C5.495 0 .16 5.335.157 11.892c0 2.096.547 4.142 1.588 5.945L.057 24l6.305-1.654a11.882 11.882 0 005.683 1.448h.005c6.554 0 11.89-5.335 11.893-11.893a11.821 11.821 0 00-3.48-8.413z"/>
        </svg>
        {post['cta']}
      </div>
      <div class="handle">@despachante_sergioneto<small>Campo do Kigol</small></div>
    </div>
  </div>
</body>
</html>
"""


def screenshot(html_path: Path, png_path: Path) -> None:
    cmd = [
        str(EDGE),
        "--headless=new",
        "--disable-gpu",
        "--hide-scrollbars",
        f"--screenshot={png_path}",
        "--window-size=1080,1080",
        html_path.as_uri(),
    ]
    subprocess.run(cmd, check=True, capture_output=True)
    for _ in range(25):
        if png_path.exists() and png_path.stat().st_size > 5000:
            return
        time.sleep(0.2)
    raise RuntimeError(f"Screenshot failed: {png_path}")


def build_docx(paths: list[tuple[dict, Path]]) -> Path:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    h = doc.add_heading("Criativos com foto — Despachante Sérgio Neto", level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x33, 0x58)

    p = doc.add_paragraph()
    r = p.add_run(
        "Upload deste DOCX no Google Drive → Abrir com Documentos Google. "
        "Publique a imagem + a copy na legenda (Google Meu Negócio e Instagram)."
    )
    r.italic = True

    doc.add_paragraph("WhatsApp: (19) 99185-1849 · (19) 99791-1339")
    doc.add_paragraph("Instagram: @despachante_sergioneto · Campo do Kigol, Cosmópolis-SP")
    doc.add_paragraph()

    for post, img in paths:
        hh = doc.add_heading(f"{post['slug'].upper()} — {post['dia']}", level=1)
        for run in hh.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x33, 0x58)

        cb = doc.add_paragraph()
        cb.add_run("☐  ").bold = True
        cb.add_run("Publicado no Google — data: __________")
        cb2 = doc.add_paragraph()
        cb2.add_run("☐  ").bold = True
        cb2.add_run("Publicado no Instagram — data: __________")

        doc.add_paragraph("Criativo:")
        doc.add_picture(str(img), width=Inches(4.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        lab = doc.add_paragraph()
        r = lab.add_run("Copy (colar na legenda):")
        r.bold = True
        doc.add_paragraph(post["copy"])
        doc.add_paragraph("— — — — — — — — — — — — — — —")

    out = OUT / "posts-criativos-despachante.docx"
    doc.save(out)
    return out


def main() -> None:
    results: list[tuple[dict, Path]] = []
    for post in POSTS:
        html_path = OUT / f"{post['slug']}.html"
        png_path = OUT / f"{post['slug']}.png"
        html_path.write_text(html_for(post), encoding="utf-8")
        print("render", post["slug"])
        screenshot(html_path, png_path)
        results.append((post, png_path))

    docx = build_docx(results)
    print("docx", docx)

    cards = []
    for post, img in results:
        cards.append(
            f"<section class='card'><h2>{post['slug']} · {post['dia']}</h2>"
            f"<img src='{img.name}' width='540' alt=''/>"
            f"<pre>{post['copy']}</pre>"
            f"<p>☐ Google &nbsp; ☐ Instagram</p></section>"
        )
    (OUT / "galeria-posts.html").write_text(
        "<!DOCTYPE html><html lang='pt-BR'><head><meta charset='UTF-8'/>"
        "<title>Galeria posts</title><style>"
        "body{font-family:Arial,sans-serif;max-width:720px;margin:2rem auto;padding:0 1rem;color:#1a3358}"
        "img{max-width:100%;border-radius:10px;box-shadow:0 10px 28px rgba(0,0,0,.15)}"
        "pre{background:#eef3f9;padding:1rem;white-space:pre-wrap;border-left:4px solid #e8a317}"
        ".card{margin:2rem 0;padding-bottom:1.5rem;border-bottom:1px solid #dce6f2}"
        "</style></head><body><h1>Criativos com contexto</h1>"
        + "".join(cards)
        + "</body></html>",
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
